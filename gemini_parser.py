import os
import io
import docx
import pypdf
import re
import time
from fishery_schema import FisheryLogBatchSchema, BiologicalParameterBatch
from typing import Tuple, Dict, Any, Optional, Union

try:
    from google import genai
    from google.genai import types
except Exception:
    genai = None
    types = None

def extract_docx_text(file_bytes: bytes) -> str:
    """Extracts both paragraph text and table data from a DOCX file."""
    doc = docx.Document(io.BytesIO(file_bytes))
    lines = []
    
    # Extract paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            lines.append(paragraph.text)
            
    # Extract tables
    for table in doc.tables:
        lines.append("\n--- Table Data ---")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            # De-duplicate adjacent identical cells (common in merged cells)
            cleaned_cells = []
            for cell in cells:
                if not cleaned_cells or cleaned_cells[-1] != cell:
                    cleaned_cells.append(cell)
            lines.append(" | ".join(cleaned_cells))
        lines.append("------------------\n")
        
    return "\n".join(lines)

def extract_pdf_text_fallback(file_bytes: bytes) -> str:
    """Extracts text from PDF as a fallback mechanism."""
    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    text = []
    for i, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text:
            text.append(f"--- Page {i+1} ---\n{page_text}")
    return "\n".join(text)


def _is_retryable_provider_error(exc: Exception) -> bool:
    message = str(exc).lower()
    retry_markers = [
        "503",
        "unavailable",
        "high demand",
        "try again later",
        "rate limit",
        "429",
        "resource exhausted",
        "timeout",
        "temporarily unavailable",
        "internal error",
        "500",
        "502",
        "504",
    ]
    return any(marker in message for marker in retry_markers)


def _call_with_retry(request_fn, max_attempts: int = 3, base_delay_seconds: float = 1.5):
    last_error = None
    for attempt in range(1, max_attempts + 1):
        try:
            return request_fn()
        except Exception as exc:
            last_error = exc
            if attempt >= max_attempts or not _is_retryable_provider_error(exc):
                raise
            time.sleep(base_delay_seconds * attempt)

    if last_error:
        raise last_error
    raise RuntimeError("Provider request failed without a captured exception.")


def _run_gemini_json_request(
    client,
    model_name: str,
    contents,
    generation_config,
    is_bio_db: bool,
    target_database_type: str,
):
    response = _call_with_retry(
        lambda: client.models.generate_content(
            model=model_name,
            contents=contents,
            config=generation_config,
        )
    )

    repaired_text = repair_truncated_json(response.text)

    import json

    try:
        parsed_dict = json.loads(repaired_text)
    except Exception as e:
        raise ValueError(f"Failed to parse repaired JSON: {e}\nRaw JSON: {repaired_text}")

    return normalize_dfis_payload(parsed_dict, is_bio_db, target_database_type)


def _merge_fishery_logs(logs):
    merged = []
    seen = {}

    for log in logs:
        if not isinstance(log, dict):
            continue

        key = (
            str(log.get("vessel_name", "")).strip(),
            str(log.get("log_date", "")).strip(),
            str(log.get("gear_type", "")).strip(),
            str(log.get("database_type", "")).strip(),
        )

        if key not in seen:
            clone = dict(log)
            clone["gear_properties"] = dict(log.get("gear_properties") or {})
            clone["catch_records"] = list(log.get("catch_records") or [])
            seen[key] = clone
            merged.append(clone)
            continue

        existing = seen[key]
        existing["gear_properties"].update(log.get("gear_properties") or {})
        existing["catch_records"].extend(log.get("catch_records") or [])

    return merged

def parse_document_with_gemini(
    file_bytes: bytes,
    file_name: str,
    mime_type: str,
    api_key: str,
    target_database_type: str,
    model_name: str = "gemini-1.5-pro"
) -> Union[FisheryLogBatchSchema, BiologicalParameterBatch]:
    """
    Parses the file contents with Gemini Structured Outputs.
    If target_database_type is '生物學參數資料庫', returns BiologicalParameterBatch.
    Otherwise, returns FisheryLogBatchSchema.
    """
    if not api_key:
        raise ValueError("Gemini API Key is missing. Please set it in the sidebar or env variables.")

    global genai, types
    if genai is None or types is None:
        try:
            from google import genai as _genai
            from google.genai import types as _types
        except ImportError as exc:
            raise RuntimeError("尚未安裝 google-genai 套件，請先安裝 requirements.txt。") from exc
        genai = _genai
        types = _types
        
    client = genai.Client(api_key=api_key)
    
    # Check if target is biological parameters database
    is_bio_db = (target_database_type == "生物學參數資料庫")
    
    import database
    
    # Query database to get standard parameters for prompt alignment
    try:
        db_vessels = database.get_vessels()["name"].tolist()
        db_ports = database.get_ports()["name"].tolist()
        db_species = database.get_species()["chinese_name"].tolist()
    except Exception:
        db_vessels, db_ports, db_species = [], [], []
        
    db_vessels_str = ", ".join(db_vessels) if db_vessels else "無限制"
    db_ports_str = ", ".join(db_ports) if db_ports else "無限制"
    db_species_str = ", ".join(db_species) if db_species else "無限制"
    
    # Configure generation parameters.
    #
    # Gemini Developer API mode does not support JSON schema fragments that
    # contain `additionalProperties`, which Pydantic emits for our flexible
    # dict fields such as `gear_properties`, `catch_properties`,
    # `extra_properties`, and `background_properties`.
    #
    # To keep DFIS downstream validation unchanged, we request JSON text from
    # Gemini and then validate it locally with the same Pydantic models.
    generation_config = {
        "response_mime_type": "application/json",
        "temperature": 0.1,
        "max_output_tokens": 16000,
    }

    if is_bio_db:
        
        prompt = (
            "你是一個專門的魚類生物學與生殖參數數據解析器。請分析上傳的檔案內容，"
            "從中自動提取、校正、翻譯所有個體測量紀錄，並完全對齊 BiologicalParameterBatch 格式，將每行魚類記錄填入 records 陣列回傳。\n\n"
            "特別注意事項（重要）：\n"
            "1. 【採集日期】collection_date 請仔細辨識報表最上方手寫的年、月、日，並統一轉換格式為 YYYY-MM-DD。例如：『2025年8月7日』應轉換為 '2025-08-07'。請勿使用預設或模糊的日期。\n"
            "2. 【採集編號與多頁完整提取】collection_id 填入個體編號。\n"
            "   - 極重要：每頁報表通常採用左右並排的雙表格版面（例如左半邊為編號 1-25，右半邊為編號 26-50，每頁共 50 筆）。\n"
            "   - 如果上傳的文件或 PDF 包含『多個頁面/分頁』，請務必『逐頁依序完整提取所有頁面』的所有個體記錄！例如，若有 2 頁，總共應提取出 100 筆記錄；若有 3 頁，應提取出 150 筆記錄。請絕對不能只提取第一頁或只提取前幾筆就停止！請將所有頁面的記錄合併在同一個 records 陣列中回傳。\n"
            "3. 【同上/垂直線向下填滿規則】極重要：在魚種欄位中，如果畫有垂直直線「|」、「‖」，或是留空，或是寫著「同上」，這代表該列個體的魚種與『上一列（上方最近的有值行）』完全相同。請務必自動將上一列的魚種名稱填入該列（向下填滿，ditto fill-down），絕對不能留空或忽略該行！\n"
            "4. 【採集港口】port 填入港口名稱。例如報表中的『作業地點：中芸』，則 port 填入 '中芸'。\n"
            "5. 【船名與表格代碼】vessel_name 填入船名，form_code 填入表單代碼，例如 1017。\n"
            "6. 【魚種名稱】species_name 填入魚種標準中文俗名，例如 大棘大眼鯛。\n"
            "7. 【性別與成熟度】sex 填入性別(雄性、雌性)；maturity 填入成熟度描述(如: 稍有精液、成熟、水卵)。若無此資料請填 null。\n"
            "8. 【全長/體長單位換算】total_length_mm 單位必須為毫米(mm)。如果表單原始數值是以公分(cm)表示，請務必乘以 10 自動換算成毫米(mm)。例如：20.84 公分 -> 填入 208.4。\n"
            "9. 【體重單位換算】weight_g 單位必須為公克(g)。如果表單原始數值是以公斤(kg)表示，請務必乘以 1000 自動換算成公克(g)。例如：0.14435 公斤 -> 填入 144.35。\n"
            "10. 【生殖腺指數】gsi 填入生殖腺量測對應的指數，若無此資料請填 null。\n"
            "11. 【備註】remarks 填入任何備註（如第幾網、作業地點備註等），若無此資料請填 null。\n"
            "12. 【標準資料庫連動參照（優先對齊）】：\n"
            f"   - 船名優先對齊：{db_vessels_str}\n"
            f"   - 港口優先對齊：{db_ports_str}\n"
            f"   - 魚種名稱優先標準化為對照代碼管理中的標準中文名：{db_species_str}\n"
            "   請務必將識別出的船名、港口、魚種與上述清單進行模糊比對，若有相近項目請優先標準化為清單中的字樣。"
        )
    else:
        prompt = (
            "你是一個專業的海洋漁業資訊系統解析器。請分析上傳的檔案內容，"
            "從中自動提取、分割、校正、除錯，並完全對齊 FisheryLogBatchSchema 的格式，將解析出的所有作業日誌以 logs 陣列回傳。\n\n"
            "特別注意事項（重要）：\n"
            "1. 【資料庫分類限制】本批次上傳的所有航次紀錄已明確指定資料庫分類類型，請將所有 logs 的 database_type 欄位值直接設定為：'" + target_database_type + "'。\n"
            "2. 【多重日期與多頁分拆】如果原始報表、圖片或 PDF 中包含多個日期、時間或不同地點的作業紀錄，或者包含『多個頁面/分頁』，"
            "請務必將其完整提取，並依照「每艘船每日作業」分別提取為獨立的 FisheryLogSchema 物件存入 logs 列表中，絕對不能漏掉任何一頁或任何一個航次！\n"
            "3. 【手寫劃除替代規則】極重要！如果手寫日誌中有將印刷字體或舊文字畫線塗改（劃除）並手寫填上新魚種的情形："
            "例如劃除 '白帶魚' 並手寫改為 '石姥'；劃除 '黃雞魚' 並手寫改為 '黃石斑'；劃除 '紅目鰱' 並手寫改為 '紅盤' 等，"
            "AI 判讀時必須完全以『劃除後新寫上去的魚種』為準，將新魚種填入原始魚種名稱 (species_raw_name)，並自動對應其標準名稱 (species_standard_name)，"
            "絕對不要填入已被劃除的舊魚種名。\n"
            "4. 【原始魚名與標準名】species_raw_name 填入手寫改寫後的原樣（如: '石姥', '黃石斑', '紅盤'），species_standard_name 請對齊臺灣魚類資料庫的標準中文俗名或學名（例如: '石姥' -> '波紋唇魚', '黃石斑' -> '青石斑魚', '紅盤' -> '真鯛'）。\n"
            "5. 【重量單位】若重量原始單位為台斤、台兩、公克(g)或其他單位，請務必自動換算為公斤(kg)再填入。1台斤 = 0.6公斤。\n"
            "6. 【動態參數】請將所有與特定漁獲單尾相關的量測數據（如尾數、體長等）放入 catch_properties 中，將漁法/作業參數（如經緯度、水深、下竿時數、出進港時間等）放入 gear_properties 中。"
        )
        
        if target_database_type == "休閒船釣漁業資料庫":
            prompt += (
                "\n\n7. 【休閒船釣漁業資料庫欄位強約束規格】：\n"
                "   - 漁法/作業參數（請提取並填入 logs 的 gear_properties 中）：\n"
                "     * `submitter`: 填表人姓名\n"
                "     * `start_time`: 作業時間 (格式如 HH:MM)\n"
                "     * `end_time`: 結束時間 (格式如 HH:MM)\n"
                "     * `start_latitude`: 作業經緯度 N 座標 (例如: 22 25.1N -> '22 25.1')\n"
                "     * `start_longitude`: 作業經緯度 E 座標 (例如: 120 22.4E -> '120 22.4')\n"
                "     * `end_latitude`: 結束經緯度 N 座標\n"
                "     * `end_longitude`: 結束經緯度 E 座標\n"
                "     * `grid_number`: 作業漁區的網格編號，為 1 到 52 之間的整數數字。如果地圖、格線上標記或圈選了對應網格編號，請務必提取填入此欄位。\n"
                "   - 作業漁具漁法 (gear_type 與 gear_properties 欄位定義)：\n"
                "     * 若勾選為「流刺網」，則 `gear_type` 設定為 '流刺網'，並在 `gear_properties` 中提取 `mesh_size` (網目尺寸)、`net_length` (網長)、`depth` (深度)。\n"
                "     * 若勾選為「曳繩釣」，則 `gear_type` 設定為 '曳繩釣'，並在 `gear_properties` 中提取 `gear_count` (漁具數) 與 `bait` (餌料，可為 假餌, 白帶魚, 硬尾, 煙仔魚)。\n"
                "     * 若勾選為「延繩釣」，則 `gear_type` 設定為 '延繩釣'，並在 `gear_properties` 中提取 `gear_count` (漁具數，筐或鉤) 與 `bait` (餌料，如 硬尾, 白帶魚, 其他)。\n"
                "     * 若勾選為「一支釣」，則 `gear_type` 設定為 '一支釣'，並在 `gear_properties` 中提取 `bait` (餌料，如 硬尾, 白帶魚, 其他)。\n"
                "   - 漁獲魚種名稱的對齊：請特別注意將以下俗名標準化：\n"
                "     * '黑口' -> '黑喉石首魚'\n"
                "     * '白口' -> '白姑魚'\n"
                "     * '馬頭魚' -> '日本馬頭魚'\n"
                "     * '紅甘' -> '高體鰤'\n"
                "     * '黃石斑' -> '青石斑魚'\n"
                "     * '赤鯮' -> '黃背牙鯛'\n"
                "     * '黃雞魚' -> '三線磯鱸'\n"
                "     * '盤仔' -> '魬鯛'\n"
                "     * '紅大目' -> '大眼鯛'\n"
                "     * '嘉鱲' -> '真鯛'\n"
                "     * '鰹魚' -> '正鰹'\n"
                "     * '土魠魚' -> '康氏馬鮫'\n"
                "     * '紅肉蒜' -> '藍點鰓棘鱸'\n"
                "     * '紅魚仔' -> '赤鰭笛鯛'\n"
                "     * '海鯽仔舅' -> '短棘鰏'\n"
                "     * '花鱸(公)' -> '花鱸'\n"
                "     * '鍵仔魚' -> '藍圓鰺'\n"
                "     * '黑大目' -> '異鱗鰒'"
            )
            
        if target_database_type == "刺網類漁業報表資料庫":
            prompt += (
                "\n\n7. 【刺網類漁業報表資料庫欄位強約束規格】：\n"
                "   - 漁法/作業參數（請提取並填入 logs 的 gear_properties 中）：\n"
                "     * `departure_time`: 出港時間 (格式如 HH:MM)\n"
                "     * `arrival_time`: 進港時間 (格式如 HH:MM)\n"
                "     * `setting_time`: 下網時間 (格式如 HH:MM)\n"
                "     * `hauling_time`: 起網時間 (格式如 HH:MM)\n"
                "     * `setting_longitude`: 下網經度 (文字)\n"
                "     * `setting_latitude`: 下網緯度 (文字)\n"
                "     * `operation_depth`: 作業深度 (文字)\n"
                "     * `gear_length`: 網具長度 (文字)\n"
                "     * `mesh_size`: 網目尺寸 (文字)\n"
                "   - 漁獲魚種名稱的對齊：請特別注意將以下俗名標準化：\n"
                "     * '白加網(白口)' 或 '白加網' 或 '白口' -> '白姑魚'\n"
                "     * '黑加網(黑口)' 或 '黑加網' 或 '黑口' -> '黑喉石首魚'\n"
                "     * '白三牙' -> '長體石首魚'\n"
                "     * '紅三牙' -> '褐石首魚'\n"
                "     * '春子' -> '叫姑魚'\n"
                "     * '紅沙' -> '黃臘鰺'\n"
                "     * '紅槽' -> '銀紋笛鯛'\n"
                "     * '銅盤仔' -> '黃背牙鯛'\n"
                "     * '烏格(黑鯛)' 或 '烏格' -> '黑鯛'\n"
                "     * '赤翅仔' -> '黃鰭棘鯛'\n"
                "     * '枋頭(黃錫鯛)' 或 '枋頭' -> '平鯛'\n"
                "     * '金目鱸' -> '尖吻鱸'\n"
                "     * '七星鱸' -> '花鱸'\n"
                "     * '嘉鱲' -> '真鯛'\n"
                "     * '紅目鰱' -> '大眼鯛'\n"
                "     * '狗母魚' -> '多齒蛇鯔'\n"
                "     * '金鐘(三角仔)' 或 '金鐘' -> '短棘鰏'\n"
                "     * '肉魚' -> '刺鯧'\n"
                "     * '赤鯮' -> '黃背牙鯛'\n"
                "     * '金線魚' -> '金線魚'\n"
                "     * '烏魚' -> '鯔魚'\n"
                "     * '紅蝦' -> '中華管鞭蝦'\n"
                "     * '遠海梭子蟹' -> '遠海梭子蟹'\n"
                "     * '花蟹(鏽斑蟳)' 或 '花蟹' -> '鏽斑蟳'\n"
                "     * '三點仔(紅星梭子蟹)' 或 '三點仔' -> '紅星梭子蟹'\n"
                "     * '帕頭' -> '黑喉石首魚'\n"
                "     * '台仔(星雞魚)' 或 '台仔' -> '星雞魚'\n"
                "     * '𩵚魠(銀雞魚)' -> '銀雞魚'\n"
                "     * '牛尾' -> '印度鯒'\n"
                "     * '青嘴' -> '單斑龍占'\n"
                "     * '雞仔魚' -> '三線磯鱸'\n"
                "     * '加志' -> '斑雞魚'\n"
                "     * '點誌' -> '單斑笛鯛'\n"
                "     * '臭肚' -> '臭肚魚'\n"
                "     * '沙腸仔' -> '多鱗鱚'\n"
                "     * '花身仔' -> '花身鯻'\n"
                "     * '小管' -> '鎖管'\n"
                "     * '花枝' -> '虎斑烏賊'\n"
                "     * '上魠' -> '康氏馬鮫'\n"
                "     * '馬頭' -> '日本馬頭魚'\n"
                "     * '鐵甲' -> '泰國鰺'\n"
                "     * '虎鰻' -> '海鰻'\n"
                "     * '扁魚' -> '牙鮃'\n"
                "     * '石斑' -> '青石斑魚'\n"
                "     * '紅魚仔' -> '赤鰭笛鯛'\n"
                "     * '三角' -> '短棘鰏'"
            )
            
        if target_database_type == "釣具類漁業報表資料庫":
            prompt += (
                "\n\n7. 【釣具類漁業報表資料庫欄位強約束規格】：\n"
                "   - 漁法/作業參數（請提取並填入 logs 的 gear_properties 中）：\n"
                "     * `vessel_owner`: 船主姓名\n"
                "     * `vessel_registration`: 船籍編號/船編\n"
                "     * `departure_time`: 出港時間 (格式如 HH:MM)\n"
                "     * `arrival_time`: 進港時間 (格式如 HH:MM)\n"
                "     * `total_fishing_hours`: 總下竿時數 (數值)\n"
                "     * `bait_type`: 釣餌類型 (文字，如 秋刀魚, 小管, 活蝦, 煙仔 等，如果報表中有註記，請務必提取)\n"
                "     * `location_1_latitude`: 地點 1 緯度 (如 `22 15.3`)\n"
                "     * `location_1_longitude`: 地點 1 經度 (如 `120 10.4`)\n"
                "     * `location_1_depth`: 地點 1 水深 (數值，單位為公尺)\n"
                "     * `location_2_latitude`: 地點 2 緯度\n"
                "     * `location_2_longitude`: 地點 2 經度\n"
                "     * `location_2_depth`: 地點 2 水深 (數值，單位為公尺)\n"
                "   - 漁獲魚種名稱的對齊：請特別注意將以下俗名標準化：\n"
                "     * '赤鯮' 或 '赤宗' -> '黃背牙鯛' (請在對應 catch_properties 字典中以 `size` 標註為 '大'、'中'、'小')\n"
                "     * '馬頭' -> '日本馬頭魚'\n"
                "     * '黑喉' -> '黑喉石首魚'\n"
                "     * '紅甘' -> '高體鰤'\n"
                "     * '白帶魚' -> '肥帶鰏'\n"
                "     * '黃雞魚' -> '三線磯鱸'\n"
                "     * '紅目鰱' -> '大眼鯛'\n"
                "     * '紅槽' -> '銀紋笛鯛'\n"
                "     * '鮸' -> '鮸魚'\n"
                "     * '紅臭魚' -> '大眼鯛'\n"
                "     * '赤筆' -> '星點笛鯛'\n"
                "     * '加誌/竹梭' 或 '加誌' -> '斑雞魚'\n"
                "     * '竹梭' -> '黃尾魣'"
            )
    
    if mime_type == "application/pdf" or file_name.lower().endswith(".pdf"):
        try:
            import pypdfium2 as pdfium

            pdf = pdfium.PdfDocument(file_bytes)
            page_count = len(pdf)
            page_results = []

            for page_index, page in enumerate(pdf, start=1):
                pil_img = page.render(scale=2).to_pil()
                img_byte_arr = io.BytesIO()
                pil_img.save(img_byte_arr, format="JPEG", quality=85)
                page_image_bytes = img_byte_arr.getvalue()

                page_contents = [
                    f"以下是 PDF 第 {page_index} 頁（共 {page_count} 頁），請只辨識本頁可見內容。",
                    types.Part.from_bytes(
                        data=page_image_bytes,
                        mime_type="image/jpeg",
                    ),
                    prompt + f"\n\n請僅回傳第 {page_index} 頁的辨識結果。若本頁沒有有效紀錄，請回傳空的 logs 或 records。",
                ]
                page_results.append(
                    _run_gemini_json_request(
                        client=client,
                        model_name=model_name,
                        contents=page_contents,
                        generation_config=generation_config,
                        is_bio_db=is_bio_db,
                        target_database_type=target_database_type,
                    )
                )

            if is_bio_db:
                parsed_dict = {
                    "records": [
                        record
                        for item in page_results
                        for record in item.get("records", [])
                    ]
                }
            else:
                parsed_dict = {
                    "logs": _merge_fishery_logs(
                        [
                            log
                            for item in page_results
                            for log in item.get("logs", [])
                        ]
                    )
                }

            parsed_dict = align_payload_dates_with_filename(
                parsed_dict,
                file_name=file_name,
                is_bio_db=is_bio_db,
            )

            if is_bio_db:
                if "records" not in parsed_dict or not isinstance(parsed_dict["records"], list):
                    parsed_dict["records"] = []
                for rec in parsed_dict["records"]:
                    if not isinstance(rec, dict):
                        continue
                    for field in ["sex", "maturity", "total_length_mm", "weight_g", "gsi", "remarks"]:
                        if field not in rec:
                            rec[field] = None
                return BiologicalParameterBatch.model_validate(parsed_dict)
            else:
                if "logs" not in parsed_dict or not isinstance(parsed_dict["logs"], list):
                    parsed_dict["logs"] = []
                for log in parsed_dict["logs"]:
                    if not isinstance(log, dict):
                        continue
                    if "gear_properties" not in log:
                        log["gear_properties"] = {}
                    if "catch_records" not in log or not isinstance(log["catch_records"], list):
                        log["catch_records"] = []
                    for catch in log["catch_records"]:
                        if not isinstance(catch, dict):
                            continue
                        for field in ["weight_kg", "count_individual"]:
                            if field not in catch:
                                catch[field] = None
                        if "catch_properties" not in catch:
                            catch["catch_properties"] = {}
                return FisheryLogBatchSchema.model_validate(parsed_dict)
        except Exception:
            pass

    contents = []
    
    # Initialize variables for cleanup
    temp_file_path = None
    st_temp_file_ref = None
    
    # Prepare input contents depending on MIME type
    if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file_name.lower().endswith(".docx"):
        text_content = extract_docx_text(file_bytes)
        contents.append(f"文件內容如下：\n\n{text_content}\n\n")
        contents.append(prompt)
    elif mime_type == "application/pdf" or file_name.lower().endswith(".pdf"):
        try:
            import pypdfium2 as pdfium
            import io
            import sys
            import traceback
            
            pdf = pdfium.PdfDocument(file_bytes)
            page_count = len(pdf)
            multi_page_instruction = (
                f"\n\n這是一份共 {page_count} 頁的 PDF。"
                "你必須逐頁完整讀取所有頁面，不可只擷取第 1 頁。"
                "若不同頁面代表不同航次、不同日期或不同作業紀錄，必須拆成多筆 logs/records 輸出。"
                "若後續頁面是同一筆紀錄的延續，必須合併回同一筆，不可遺漏後續頁面的欄位或漁獲明細。"
            )
            for page_index, page in enumerate(pdf, start=1):
                # Render to PIL image at scale=2 for crisp OCR resolution
                pil_img = page.render(scale=2).to_pil()
                img_byte_arr = io.BytesIO()
                pil_img.save(img_byte_arr, format='JPEG', quality=85)
                page_image_bytes = img_byte_arr.getvalue()

                contents.append(f"以下是 PDF 第 {page_index} 頁（共 {page_count} 頁）。請保留頁序並持續整合後續頁面資訊。")
                contents.append(
                    types.Part.from_bytes(
                        data=page_image_bytes,
                        mime_type="image/jpeg",
                    )
                )
            contents.append(prompt + multi_page_instruction)
        except Exception as e:
            import sys
            import traceback
            print(f"Warning: pypdfium2 rendering failed, falling back to text extraction: {e}", file=sys.stderr)
            traceback.print_exc(file=sys.stderr)
            
            text_content = extract_pdf_text_fallback(file_bytes)
            contents.append(f"PDF文字提取內容如下（備用方式）：\n\n{text_content}\n\n")
            contents.append(prompt)
    elif mime_type.startswith("image/") or file_name.lower().endswith((".png", ".jpg", ".jpeg", ".webp")):
        image_part = types.Part.from_bytes(
            data=file_bytes,
            mime_type=mime_type if mime_type.startswith("image/") else "image/jpeg",
        )
        contents.append(image_part)
        contents.append(prompt)
    else:
        try:
            text_content = file_bytes.decode("utf-8", errors="ignore")
            contents.append(f"文件內容如下：\n\n{text_content}\n\n")
            contents.append(prompt)
        except Exception as e:
            raise ValueError(f"不支援的檔案格式，且無法以文字方式讀取: {file_name}")

    # Generate content using Gemini
    response = None
    try:
        response = _call_with_retry(
            lambda: client.models.generate_content(
                model=model_name,
                contents=contents,
                config=generation_config,
            )
        )
    except Exception as e:
        message = str(e)
        if "additionalProperties is only supported in Gemini Enterprise Agent Platform mode" in message:
            raise RuntimeError(
                "Gemini Developer API 不支援目前的回應 schema 限制。"
                "系統已切換為 JSON 相容模式，但本次請求仍被舊設定攔截；"
                "請重新整理頁面後再測試一次。"
            ) from e
        raise
    finally:
        # Clean up Gemini File API reference and temp file
        if st_temp_file_ref:
            try:
                client.files.delete(name=st_temp_file_ref.name)
            except Exception:
                pass
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.remove(temp_file_path)
            except Exception:
                pass
    
    # Repair potentially truncated JSON response
    repaired_text = repair_truncated_json(response.text)
    
    # Return validated Pydantic object based on type
    import json
    try:
        parsed_dict = json.loads(repaired_text)
    except Exception as e:
        raise ValueError(f"Failed to parse repaired JSON: {e}\nRaw JSON: {repaired_text}")

    parsed_dict = normalize_dfis_payload(parsed_dict, is_bio_db, target_database_type)
    parsed_dict = align_payload_dates_with_filename(
        parsed_dict,
        file_name=file_name,
        is_bio_db=is_bio_db,
    )
        
    if is_bio_db:
        # Pre-populate missing optional fields in BiologicalParameterRecord to prevent Pydantic validation errors
        if "records" not in parsed_dict or not isinstance(parsed_dict["records"], list):
            parsed_dict["records"] = []
            
        for rec in parsed_dict["records"]:
            if not isinstance(rec, dict):
                continue
            for field in ["sex", "maturity", "total_length_mm", "weight_g", "gsi", "remarks"]:
                if field not in rec:
                    rec[field] = None
                    
        return BiologicalParameterBatch.model_validate(parsed_dict)
    else:
        # Pre-populate missing optional fields in CatchDetail and FisheryLogSchema to prevent Pydantic validation errors
        if "logs" not in parsed_dict or not isinstance(parsed_dict["logs"], list):
            parsed_dict["logs"] = []
            
        for log in parsed_dict["logs"]:
            if not isinstance(log, dict):
                continue
            if "gear_properties" not in log:
                log["gear_properties"] = {}
            if "catch_records" not in log or not isinstance(log["catch_records"], list):
                log["catch_records"] = []
                
            for catch in log["catch_records"]:
                if not isinstance(catch, dict):
                    continue
                for field in ["weight_kg", "count_individual"]:
                    if field not in catch:
                        catch[field] = None
                if "catch_properties" not in catch:
                    catch["catch_properties"] = {}
                    
        return FisheryLogBatchSchema.model_validate(parsed_dict)


def _is_missing_or_zero(value) -> bool:
    if value is None:
        return True
    if isinstance(value, (int, float)):
        return value == 0
    if isinstance(value, str):
        stripped = value.strip()
        return stripped in {"", "0", "0.0", "0.000", "null", "None"}
    return False


def _coerce_json_dict(value):
    if isinstance(value, dict):
        return value
    if isinstance(value, str) and value.strip():
        try:
            import json
            parsed = json.loads(value)
            return parsed if isinstance(parsed, dict) else {}
        except Exception:
            return {}
    return {}


def infer_document_year(file_name: str) -> Optional[int]:
    """Infer the expected Gregorian year from the uploaded filename."""
    if not file_name:
        return None

    gregorian_matches = re.findall(r"(?<!\d)(20\d{2})(?!\d)", file_name)
    if gregorian_matches:
        year = int(gregorian_matches[0])
        if 2000 <= year <= 2099:
            return year

    roc_matches = re.findall(r"(?<!\d)(1\d{2})(?!\d)", file_name)
    for raw in roc_matches:
        roc_year = int(raw)
        if 100 <= roc_year <= 199:
            return roc_year + 1911

    return None


def _rewrite_date_year(date_value, expected_year: Optional[int]):
    if not expected_year or not isinstance(date_value, str):
        return date_value

    stripped = date_value.strip()
    match = re.match(r"^(\d{4})-(\d{2})-(\d{2})$", stripped)
    if not match:
        return date_value

    current_year = int(match.group(1))
    month = int(match.group(2))
    day = int(match.group(3))
    if not (1 <= month <= 12 and 1 <= day <= 31):
        return date_value

    if current_year == expected_year:
        return stripped

    if 2000 <= current_year <= 2099:
        return f"{expected_year:04d}-{month:02d}-{day:02d}"

    return date_value


def align_payload_dates_with_filename(
    parsed_dict,
    file_name: str,
    is_bio_db: bool,
):
    """Correct OCR-misread years using the document filename as a stable hint."""
    expected_year = infer_document_year(file_name)
    if not expected_year or not isinstance(parsed_dict, dict):
        return parsed_dict

    if is_bio_db:
        records = parsed_dict.get("records")
        if isinstance(records, list):
            for record in records:
                if not isinstance(record, dict):
                    continue
                if "collection_date" in record:
                    record["collection_date"] = _rewrite_date_year(
                        record.get("collection_date"),
                        expected_year,
                    )
        return parsed_dict

    logs = parsed_dict.get("logs")
    if isinstance(logs, list):
        for log in logs:
            if not isinstance(log, dict):
                continue
            if "log_date" in log:
                log["log_date"] = _rewrite_date_year(
                    log.get("log_date"),
                    expected_year,
                )
    return parsed_dict


def _stitch_fragmented_fishery_logs(logs):
    """Fill missing vessel/date context for continuation pages in multi-page PDFs."""
    stitched = []
    last_context = None

    for log in logs:
        if not isinstance(log, dict):
            continue

        current = dict(log)
        current["gear_properties"] = dict(current.get("gear_properties") or {})
        current["catch_records"] = list(current.get("catch_records") or [])

        has_vessel = bool(str(current.get("vessel_name", "")).strip())
        has_log_date = bool(str(current.get("log_date", "")).strip())
        has_catches = len(current["catch_records"]) > 0
        has_gear_type = bool(str(current.get("gear_type", "")).strip())

        if last_context and has_catches:
            if not has_vessel and last_context.get("vessel_name"):
                current["vessel_name"] = last_context["vessel_name"]
                has_vessel = True
            if not has_log_date and last_context.get("log_date"):
                current["log_date"] = last_context["log_date"]
                has_log_date = True
            if not has_gear_type and last_context.get("gear_type"):
                current["gear_type"] = last_context["gear_type"]
            if not current.get("database_type") and last_context.get("database_type"):
                current["database_type"] = last_context["database_type"]
            if not current["gear_properties"] and last_context.get("gear_properties"):
                current["gear_properties"] = dict(last_context["gear_properties"])

        if has_vessel or has_log_date:
            last_context = {
                "vessel_name": current.get("vessel_name"),
                "log_date": current.get("log_date"),
                "gear_type": current.get("gear_type"),
                "database_type": current.get("database_type"),
                "gear_properties": dict(current.get("gear_properties") or {}),
            }

        stitched.append(current)

    return stitched


def normalize_dfis_payload(parsed_dict, is_bio_db: bool, target_database_type: str):
    """Normalize common model output variants into DFIS canonical payloads."""
    if is_bio_db:
        if isinstance(parsed_dict, list):
            return {"records": parsed_dict}
        if isinstance(parsed_dict, dict) and "records" not in parsed_dict:
            if any(key in parsed_dict for key in ["collection_date", "species_name", "form_code"]):
                return {"records": [parsed_dict]}
        return parsed_dict

    if isinstance(parsed_dict, list):
        parsed_dict = {"logs": parsed_dict}
    elif isinstance(parsed_dict, dict) and "logs" not in parsed_dict:
        if any(
            key in parsed_dict
            for key in [
                "database_type",
                "boat_name",
                "vessel_name",
                "date",
                "log_date",
                "catches",
                "catch_records",
            ]
        ):
            parsed_dict = {"logs": [parsed_dict]}

    if not isinstance(parsed_dict, dict):
        return parsed_dict

    logs = parsed_dict.get("logs")
    if not isinstance(logs, list):
        return parsed_dict

    candidate_logs = []
    for log in logs:
        if not isinstance(log, dict):
            continue

        nested_data = log.get("data")
        if isinstance(nested_data, dict):
            merged_log = dict(nested_data)
            for key, value in log.items():
                if key == "data":
                    continue
                if key not in merged_log or merged_log.get(key) in [None, "", [], {}]:
                    merged_log[key] = value
            log.clear()
            log.update(merged_log)

        log.setdefault("database_type", target_database_type)
        if "ship" in log and "vessel_name" not in log:
            log["vessel_name"] = log.pop("ship")
        if "ship_name" in log and "vessel_name" not in log:
            log["vessel_name"] = log.pop("ship_name")
        if "boat_name" in log and "vessel_name" not in log:
            log["vessel_name"] = log.pop("boat_name")
        if "date" in log and "log_date" not in log:
            log["log_date"] = log.pop("date")
        if "work_date" in log and "log_date" not in log:
            log["log_date"] = log.pop("work_date")
        if "day" in log and "log_date" not in log:
            log["log_date"] = log.pop("day")
        if "fishing_method" in log and "gear_type" not in log:
            log["gear_type"] = log.pop("fishing_method")
        if "submitter" in log and "observer_name" not in log:
            log["observer_name"] = log["submitter"]

        catch_records = log.get("catch_records")
        if catch_records is None and "catches" in log:
            catch_records = log.pop("catches")
            log["catch_records"] = catch_records
        if catch_records is None and "catch" in log:
            catch_records = log.pop("catch")
            log["catch_records"] = catch_records
        if catch_records is None and "records" in log and isinstance(log.get("records"), list):
            catch_records = log.pop("records")
            log["catch_records"] = catch_records

        if catch_records is None and any(
            key in log
            for key in [
                "original_name",
                "species_raw_name",
                "standard_name",
                "species_standard_name",
                "weight_kg",
                "count",
                "count_individual",
                "catch_properties",
            ]
        ):
            single_catch = {}
            for key in [
                "original_name",
                "species_raw_name",
                "standard_name",
                "species_standard_name",
                "translated_name",
                "species_name",
                "weight_kg",
                "count",
                "count_individual",
                "catch_properties",
            ]:
                if key in log:
                    single_catch[key] = log.pop(key)
            log["catch_records"] = [single_catch]

        gear_props = log.get("gear_properties")
        if not isinstance(gear_props, dict):
            gear_props = {}
            log["gear_properties"] = gear_props

        for source_key in [
            "submitter",
            "start_time",
            "end_time",
            "start_latitude",
            "start_longitude",
            "end_latitude",
            "end_longitude",
            "bait",
            "number",
        ]:
            if source_key in log and source_key not in gear_props:
                gear_props[source_key] = log[source_key]

        if "gear_type" not in log or not log.get("gear_type"):
            log["gear_type"] = gear_props.get("gear_type") or target_database_type

        if isinstance(log.get("catch_records"), list):
            for catch in log["catch_records"]:
                if not isinstance(catch, dict):
                    continue
                if "original_name" in catch and "species_raw_name" not in catch:
                    catch["species_raw_name"] = catch.pop("original_name")
                if "standard_name" in catch and "species_standard_name" not in catch:
                    catch["species_standard_name"] = catch.pop("standard_name")
                if "translated_name" in catch and "species_standard_name" not in catch:
                    catch["species_standard_name"] = catch.pop("translated_name")
                if "species_name" in catch and "species_standard_name" not in catch:
                    catch["species_standard_name"] = catch["species_name"]
                if "count" in catch and "count_individual" not in catch:
                    catch["count_individual"] = catch["count"]
                catch["catch_properties"] = _coerce_json_dict(catch.get("catch_properties"))
                catch_props = catch["catch_properties"]
                if isinstance(catch_props, dict):
                    if (
                        ("count_individual" not in catch or _is_missing_or_zero(catch["count_individual"]))
                        and "count" in catch_props
                    ):
                        catch["count_individual"] = catch_props["count"]
                    if (
                        ("weight_kg" not in catch or _is_missing_or_zero(catch["weight_kg"]))
                        and "weight_kg" in catch_props
                    ):
                        catch["weight_kg"] = catch_props["weight_kg"]

        candidate_logs.append(log)

    stitched_logs = _stitch_fragmented_fishery_logs(candidate_logs)
    normalized_logs = []
    for log in stitched_logs:
        has_vessel = bool(str(log.get("vessel_name", "")).strip())
        has_log_date = bool(str(log.get("log_date", "")).strip())
        has_catches = isinstance(log.get("catch_records"), list) and len(log.get("catch_records")) > 0
        has_meaningful_gear_props = isinstance(log.get("gear_properties"), dict) and len(log.get("gear_properties")) > 0

        if (has_vessel and has_log_date) or (has_log_date and has_catches) or (has_vessel and has_catches):
            normalized_logs.append(log)
        elif has_vessel and has_meaningful_gear_props:
            normalized_logs.append(log)

    parsed_dict["logs"] = normalized_logs
    return parsed_dict

def repair_truncated_json(json_str: str) -> str:
    """
    Attempts to repair a truncated JSON string returned by Gemini.
    Looks backward for the last complete object bracket '}' and appends
    necessary closing tags (e.g. '] }') to make it parseable.
    """
    import json

    json_str = json_str.strip()
    if json_str.startswith("```"):
        json_str = json_str.strip("`")
        if json_str.lower().startswith("json"):
            json_str = json_str[4:].strip()

    first_object = min(
        [idx for idx in [json_str.find("{"), json_str.find("[")] if idx != -1],
        default=-1,
    )
    if first_object > 0:
        json_str = json_str[first_object:]

    try:
        json.loads(json_str)
        return json_str
    except json.JSONDecodeError:
        pass

    stack = []
    in_string = False
    escape = False
    last_safe_cut = -1

    for idx, ch in enumerate(json_str):
        if in_string:
            if escape:
                escape = False
            elif ch == "\\":
                escape = True
            elif ch == '"':
                in_string = False
            continue

        if ch == '"':
            in_string = True
        elif ch in "{[":
            stack.append(ch)
        elif ch == "}" and stack and stack[-1] == "{":
            stack.pop()
            last_safe_cut = idx + 1
        elif ch == "]" and stack and stack[-1] == "[":
            stack.pop()
            last_safe_cut = idx + 1
        elif ch == ",":
            last_safe_cut = idx

    candidate = json_str
    if in_string or stack:
        if last_safe_cut > 0:
            candidate = json_str[:last_safe_cut].rstrip(", \n\r\t")
        else:
            candidate = json_str

    stack = []
    in_string = False
    escape = False
    for ch in candidate:
        if in_string:
            if escape:
                escape = False
            elif ch == "\\":
                escape = True
            elif ch == '"':
                in_string = False
            continue
        if ch == '"':
            in_string = True
        elif ch in "{[":
            stack.append(ch)
        elif ch == "}" and stack and stack[-1] == "{":
            stack.pop()
        elif ch == "]" and stack and stack[-1] == "[":
            stack.pop()

    closing = []
    for opener in reversed(stack):
        closing.append("}" if opener == "{" else "]")

    repaired = candidate + "".join(closing)
    try:
        json.loads(repaired)
        return repaired
    except json.JSONDecodeError:
        return json_str
